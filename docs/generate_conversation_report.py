from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).with_name("Mamba_3v3_完整对话与开发报告.docx")
BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
INK = "243447"
MUTED = "667085"
FONT = "Microsoft YaHei"


def font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = tc.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tc.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, LIGHT); margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h)); font(r, 9.5, True, DARK)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]; margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: c.width = Inches(widths[i])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value)); font(r, 9.2)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            if widths: c.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.38 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), 10.5)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), 10.5)
    return p


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), 10.5, True)
        font(p.add_run(text[len(bold_lead):]), 10.5)
    else: font(p.add_run(text), 10.5)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), {1:16,2:13,3:11.5}[level], True, BLUE if level < 3 else DARK)
    return p


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT; normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT); normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for n, sz, before, after, color in ((1,16,18,10,BLUE),(2,13,14,7,BLUE),(3,11.5,10,5,DARK)):
    s = styles[f"Heading {n}"]; s.font.name=FONT; s._element.rPr.rFonts.set(qn("w:eastAsia"),FONT)
    s.font.size=Pt(sz); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run("Mamba 3v3 | 对话、规则与开发归档"), 8.5, False, MUTED)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("Booster Studio 项目记录 | 2026-07-12"), 8.5, False, MUTED)

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(120); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("项目完整归档"), 12, True, BLUE)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
font(p.add_run("Mamba 3v3 足球 Agent"), 28, True, DARK)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(28)
font(p.add_run("全部对话信息、赛事规则、代码改造、调参与提交指南"), 14, False, MUTED)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("工作目录：D:\\BoosterStudio\\mamba"), 10.5, False, INK)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("当前交付版本：com.example.mamba 2.2.0"), 10.5, True, BLUE)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80)
font(p.add_run("整理日期：2026 年 7 月 12 日"), 10, False, MUTED)
doc.add_page_break()

heading(doc, "文档说明", 1)
para(doc, "本报告完整整理了用户与 Codex 围绕 Booster 3v3 足球比赛的全部实质信息。内容包含飞书资料读取、赛事规则、飞书 CLI 安装、基础代码分析、官方模板迁移、规则修复、速度调优、自适应战术、测试、构建、部署、运行和提交方法。")
para(doc, "为保护账号安全，报告不记录已经使用过的 OAuth 设备码、用户码、二维码、App Secret、access token 或其他认证材料。")

heading(doc, "目录", 1)
for x in ["1. 需求与目标", "2. 资料来源", "3. 飞书 CLI 安装与授权", "4. 赛事规则完整摘要", "5. 官方机器人构建流程", "6. 初始代码分析", "7. 代码整体重构", "8. 自适应战术与参数", "9. 测试、构建和版本记录", "10. 运行与提交操作", "11. 风险、限制与后续建议", "12. 对话时间线"]:
    bullet(doc, x)

heading(doc, "1. 需求与目标", 1)
para(doc, "用户希望从指定飞书知识库中理解比赛机器人构建流程和方法，完整遵守赛事规则，分析本地 Agent 基础代码，随后直接完成代码改造、提升线上仿真比赛性能、生成可提交的 .agent 文件，并最终形成完整归档文档。")
bullet(doc, "比赛类型：线上 3v3 机器人足球仿真赛。")
bullet(doc, "机器人平台：Booster K1 / T1；项目最终面向 football3v3 场景。")
bullet(doc, "提交格式：.agent，而不是源码压缩包。")
bullet(doc, "最终目标：在合规前提下尽可能提高速度、决策质量、传射成功率和防守能力。")

heading(doc, "2. 资料来源", 1)
table(doc, ["资料", "链接", "用途"], [
    ("角色分配", "LY2iwt6pXi28InkVhnmcBgxQnQb", "Playbook、RoleAssignment、角色子树"),
    ("第一个 Booster Agent", "Vho8w1ylvisGYIkh9MAc7NQqnrh", "创建、构建、部署、运行流程"),
    ("选择射门队员", "G38fwu8NPiiqyzkwA14cLUq0npf", "Chaser 三层选择与出球决策"),
    ("优化移动行为", "AI5Qwwnrli5V97kCP2CcuSF8nRc", "绕行、速度控制、近邻偏航"),
    ("配置参数指南", "TOz6wp4tdiMNBMkg3oMcaE2pnCc", "四类实战参数策略"),
    ("赛场规则", "V5oAwY9QjierOYk6PkncdilUnhe", "状态、定位球、站位、处罚和合规"),
], [1.45, 2.45, 2.6])
para(doc, "完整 URL 统一为 https://booster.feishu.cn/wiki/<token>。此外参考了本机 Booster Studio 自带 resources/soccer_sim 官方模板，以及 BoosterRobotics/robocup_demo 官方开源仓库。")

heading(doc, "3. 飞书 CLI 安装与授权", 1)
para(doc, "最初网页抓取失败，因为知识库受租户权限保护，from=from_copylink 只表示链接来源，不改变访问权限。随后按飞书官方 CLI 指南安装并授权。")
table(doc, ["项目", "结果"], [
    ("Node.js", "通过 winget 安装 LTS v24.18.0"),
    ("npm", "11.16.0"),
    ("飞书 CLI", "@larksuite/cli 1.0.68"),
    ("飞书 Skills", "通过 npx skills add 安装 27 个官方 Skills"),
    ("身份", "机器人身份 ready；用户身份 ready"),
    ("文档访问", "用户授权后可读取 Wiki、Docx、Drive 等推荐权限"),
], [1.8, 4.7])
for s in ["npm PowerShell 入口受执行策略限制，因此使用 npm.cmd / lark-cli.cmd。", "通过 lark-cli config init --new 创建应用配置。", "通过 device flow 完成浏览器授权，再由 Agent 执行 --device-code 收尾。", "最终 lark-cli auth status 显示用户令牌有效。"]: bullet(doc, s)

heading(doc, "4. 赛事规则完整摘要", 1)
heading(doc, "4.1 技术合规", 2)
bullet(doc, "允许：机器人自身传感器、官方 GameControl、公开 SDK/ROS topic/Agent API、进程内计算和公开赛事常量。")
bullet(doc, "禁止：仿真内部 WebSocket/HTTP/Unix socket/共享内存、移动球或机器人、控制裁判管理接口、端口扫描、读取宿主机或其他进程的文件/变量/token/日志、读取 /proc 或 Docker 元数据、伪造或注入非公开 ROS/DDS/UDP 数据、干扰对手或主办方服务。")
bullet(doc, "技术违规可取消资格或判负；最终解释权归赛事组。")

heading(doc, "4.2 场地参数", 2)
table(doc, ["项目", "规则值"], [
    ("场地", "14.0 m × 9.0 m"), ("球门", "宽 2.6 m，高 1.8 m，深 0.6 m"),
    ("中圈", "半径 1.5 m"), ("点球区", "3.0 m × 6.0 m"),
    ("球门区", "1.0 m × 4.0 m"), ("足球", "半径 0.11 m，整球越线才出界/进球"),
    ("线宽", "0.05 m"), ("比赛时长", "仿真时间，当前常规为 10 分钟"),
], [2.2, 4.3])
para(doc, "调参文档中的点球区 2×5 m、球门区 1×3 m 属于旧值；代码以赛事规则 3×6 m 和 1×4 m 为准。")

heading(doc, "4.3 状态与时间", 2)
table(doc, ["状态/事件", "行为与时限"], [
    ("主状态", "INITIAL → READY → SET → PLAYING → FINISHED"),
    ("READY 稳定", "全部未罚机器人连续稳定 5 秒可提前进入 SET"),
    ("READY 超时", "45 秒进入 SET"), ("SET", "约 5 秒；机器人必须静止"),
    ("Kick-off 独占", "10 秒；超时 ball_free，双方均可直接进球"),
    ("其他定位球", "Throw-in / Goal kick / Corner kick 各 45 秒"),
    ("Stop", "出界、摆球等情况下所有机器人停止"),
    ("比赛结束", "时间到或手动 end 后 FINISHED"),
], [2.1, 4.4])

heading(doc, "4.4 进球、重启与站位", 2)
bullet(doc, "有效进球要求整球越过门线，位于两门柱之间、横梁下和门深内。")
bullet(doc, "Kick-off 直接进球：开球方需至少两次稳定触球，或等待 set play 过期开放。否则不计分并重开。")
bullet(doc, "定位球主罚方一触乌龙且未过期：不计分，对方角球。")
bullet(doc, "边线出界：最后触球方的对方获得 Throw-in。底线出界根据最后触球方判 Goal kick 或 Corner kick。")
bullet(doc, "防守方在主罚方触球前先触球：罚下 30 秒并重踢。")
bullet(doc, "对方定位球时距离球不得小于 1.45 m；代码使用 1.60 m 安全值。")
bullet(doc, "Kick-off：开球方可在本方半场或中圈；非开球方必须在本方半场且中圈外。")
bullet(doc, "30 秒无人触球触发坠球，READY→SET→PLAYING，双方均可抢球并直接进球。")

heading(doc, "4.5 运动与处罚", 2)
table(doc, ["规则", "阈值/处罚"], [
    ("SET 移动", "宽限后 1 秒路径 >0.15 m：罚下 30 秒"),
    ("Stop 移动", "宽限后 2 秒路径 >0.3 m：罚下 30 秒"),
    ("倒地", "upDot≤0.45 持续 10 秒：罚下 30 秒"),
    ("无活动", "PLAYING、距球≤3 m、10 秒几乎不动：罚下 30 秒"),
    ("罚时", "统一 30 秒；提前返场会重置罚时并增加 warning"),
    ("牌", "2 warnings→1 caution；2 cautions→sent off"),
], [2.0, 4.5])

heading(doc, "5. 官方机器人构建流程", 1)
for s in [
    "在 Booster Studio 创建 Python Agent，选择 K1/T1、football3v3 场景。",
    "从官方 /teamN/... ground truth topic 和 /soccer/game_controller 构造 PlayContext。",
    "行为树每帧先更新数据，再执行安全门、READY 或 PLAYING 子树。",
    "Playbook.assign_roles 动态产生 RoleAssignment。",
    "Chaser 执行脱困、定位球触球、传球、射门、带球、解围；Supporter 接应；Goalkeeper 守门；Defender 拦截。",
    "MotionController 依次进行路径绕行、速度计算和近邻偏航避让。",
    "TeamRobotManager 统一使用公开 BoosterRobot/SoccerKickManager 接口下发命令。",
    "构建生成签名 .agent，部署到仿真，完整跑场后提交。",
]: number(doc, s)

heading(doc, "6. 初始代码分析", 1)
para(doc, "原 mamba 项目只有 src/main.py、agent.toml、build.toml 和按钮资源，是官方 ExampleAgent 挥手示例。")
bullet(doc, "AgentBase 生命周期可用；BoosterRobot 已实例化；enable_auto_getup=True。")
bullet(doc, "两个 UI 按钮：Wave 和 Custom Action；Custom Action 仅写日志。")
bullet(doc, "Wave 使用 hand_wave 动作，并可取消 active task。")
bullet(doc, "不存在 GameControl、比赛状态、世界模型、角色分配、战术、避障、处罚和规则测试。")
bullet(doc, "中文字符串存在乱码。")
para(doc, "因此原代码适合作为生命周期和打包骨架，但不能直接参赛。")

heading(doc, "7. 代码整体重构", 1)
para(doc, "本机 resources/soccer_sim 中存在与飞书教程配套的官方模板。项目以该模板为基线机械同步，并保留 Agent ID com.example.mamba。")
table(doc, ["模块", "职责"], [
    ("src/main.py", "SoccerSimAgent 生命周期入口，异步启动/停止 Runtime"),
    ("src/runtime.py", "组装 SoccerKit、ROS Adapter、行为树和 30Hz 控制循环"),
    ("soccer_framework", "配置、数据类型、GameControl、ROS truth、机器人控制、日志"),
    ("behavior_tree", "数据层、安全门、READY/PLAYING、命令提交"),
    ("play", "Playbook、角色注册、Chaser/Supporter/Defender/Goalkeeper"),
    ("tactics", "几何、目标、传射、带球、接应、重启、避障、运动和踢球迟滞"),
], [1.7, 4.8])
heading(doc, "7.1 额外规则修复", 2)
bullet(doc, "角色分配先过滤 GameControl 中受罚、红牌、替补和仍在罚时的机器人。")
bullet(doc, "受罚机器人不再占用唯一 Chaser。")
bullet(doc, "固定门将不可用时，由有效机器人中位置最靠后的球员临时接替。")
bullet(doc, "未分配机器人映射为 none，并由安全层停止。")
bullet(doc, "场景切换为 football3v3。")

heading(doc, "8. 自适应战术与参数", 1)
para(doc, "早期 2.1.0 版本曾将踢球力度设为接口极限 10.0。重新精读调参文档后发现比赛建议范围为 0.8–2.5，因此 2.2.0 改为按场景使用 2.2–2.5，减少出界和无法接应。")
table(doc, ["模式", "触发", "核心效果"], [
    ("Fluid", "中场常态", "1.2m/s、1.5rad/s、小绕行半径、稳定 Chaser、顺畅直走"),
    ("Precision", "球 x≥3.5；x<2.5 退出", "传球净空1.0、传球阈值0.65、射门通道0.65、向中路带球"),
    ("Defensive", "球 x≤-3.5；或最后150秒领先", "真实 Defender、大避障半径、保守门将、强解围"),
    ("Aggressive", "最后150秒落后", "1.2m/s、1.5rad/s、带球2.0m、近接应、低传球阈值、快速压上"),
], [1.1, 2.1, 3.3])
heading(doc, "8.1 自动切换优先级", 2)
for s in ["最后 150 秒落后：Aggressive。", "最后 150 秒领先：Defensive。", "否则球在己方危险区：Defensive。", "否则球在对方前场：Precision。", "其他：Fluid。"]: number(doc, s)
para(doc, "球场区域切换使用 1 m 迟滞区，避免球在临界线附近时逐帧切换。")
heading(doc, "8.2 新增功能", 2)
bullet(doc, "shot_lane_min_score 从硬编码 0.55 改成动态参数。")
bullet(doc, "KickHysteresis 支持 profile 切换时更新 enter/exit/delay，并清除旧状态。")
bullet(doc, "Defender 站在球和己方球门之间，横向跟球、保持后场层次并面向球。")
bullet(doc, "最大线速度使用文档上限 1.2 m/s；最大角速度使用文档上限 1.5 rad/s。")

heading(doc, "9. 测试、构建和版本记录", 1)
table(doc, ["版本", "内容", "状态"], [
    ("1.0.0", "原挥手 ExampleAgent", "已过时"),
    ("2.0.0", "官方 SoccerSim 全栈 + 受罚角色修复", "构建签名成功"),
    ("2.1.0", "静态高速配置；力度曾为10.0", "已被2.2.0取代"),
    ("2.2.0", "四档自适应策略、Defender、动态射门和迟滞", "当前提交版本"),
], [1.0, 4.3, 1.2])
para(doc, "当前自动化测试共 8 项，全部通过：赛事场地尺寸、参数边界、落后激进切换、领先防守切换、Defender 分配、前场精准切换、受罚角色过滤、定位球距离。全部 Python 文件也通过 compileall。")
para(doc, "Booster Studio 官方构建系统完成 ROS2 package 构建、py_trees x86_64/aarch64 依赖打包、Agent image 组装和签名。")

heading(doc, "10. 运行与提交操作", 1)
heading(doc, "10.1 运行", 2)
for s in [
    "在 Booster Studio 打开 D:\\BoosterStudio\\mamba。",
    "启动或重启虚拟机器人，确认场景 football3v3 和机器人在线。",
    "如果出现 SSH service not connected，关闭虚拟机器人后重新启动。",
    "点击顶部‘激活、构建、部署和运行代理’按钮。",
    "确认日志出现 SoccerSimAgent activated 和 SoccerTeamRuntime started。",
    "确认 Agent 信息显示版本 2.2.0。",
]: number(doc, s)
heading(doc, "10.2 提交", 2)
para(doc, "赛事提交格式是 .agent。当前应提交：D:\\BoosterStudio\\mamba\\build\\com.example.mamba-2.2.0.agent。不要提交 2.0.0 或 2.1.0。")
para(doc, "新旧版本使用相同 ID com.example.mamba，部署新版会覆盖旧版，无需同时运行多个 Agent。")

heading(doc, "11. 风险、限制与后续建议", 1)
bullet(doc, "构建和签名已经成功，但自动部署曾因虚拟机器人 SSH 未连接而失败；需要在 Studio 中重新连接后完成整场仿真验收。")
bullet(doc, "必须分别验证 team 1 和 team 2 的环境变量和坐标方向。")
bullet(doc, "需要覆盖 READY、SET、Stop、所有定位球、受罚、倒地、少人、落后、领先、前场和后场场景。")
bullet(doc, "最大参数并不等于最大胜率。2.2.0 已按文档限制将速度拉高，同时保留射门质量和出界控制。")
bullet(doc, "RosTruthProvider 只能使用赛事明确公开的 /teamN/... topic；不得扩展到未公开仿真内部接口。")
bullet(doc, "完成仿真后如有抖动、绕远或出界，应基于结构化日志定向调参，而不是同时改变所有变量。")

heading(doc, "12. 对话时间线", 1)
timeline = [
    ("资料请求", "用户提供 5 个构建资料链接和 1 个赛事规则链接，要求只读分析本地代码。"),
    ("访问障碍", "公开网页抓取失败；解释 from_copylink 不改变权限，建议公开分享或导出。"),
    ("CLI 安装", "按官方指南安装 Node、@larksuite/cli 和 27 个 Skills。"),
    ("应用配置", "生成飞书官方配置链接，用户完成配置。"),
    ("用户授权", "完成 OAuth device flow；用户和机器人身份均 ready。"),
    ("读取资料", "成功读取全部 6 个 Wiki；发现官方教程与当前 mamba 代码不匹配。"),
    ("规则解析", "整理状态机、定位球、进球、站位、运动、处罚和技术合规。"),
    ("模板发现", "在 Booster Studio resources/soccer_sim 找到官方完整 Python 模板。"),
    ("2.0.0", "整体迁移官方模板，修复受罚角色和动态门将，添加测试并构建。"),
    ("提交确认", "用户确认提交格式为 .agent；指出应先完成端到端仿真。"),
    ("速度诊断", "解释 0.8m/s 官方默认偏保守、0.5rad 转向阈值导致先转后走。"),
    ("2.1.0", "线上仿真高速参数：1.2m/s、力度10.0等；构建成功。"),
    ("文档复读", "重新精读配置参数指南，确认角速度示例上限1.5、建议踢球上限2.5。"),
    ("2.2.0", "实现 Fluid/Precision/Defensive/Aggressive、自适应切换、Defender、动态射门阈值和迟滞。"),
    ("当前", "8 项测试通过，2.2.0 构建签名成功，等待 Studio 中整场仿真和提交。"),
]
table(doc, ["阶段", "结论/操作"], timeline, [1.25, 5.25])

heading(doc, "附录 A：关键路径", 1)
for s in [
    "Agent 入口：mamba/src/main.py",
    "运行时：mamba/src/runtime.py",
    "配置：mamba/src/soccer_framework/config.py",
    "赛事类型：mamba/src/soccer_framework/types.py",
    "角色：mamba/src/play/playbook.py、default_roles.py",
    "自适应策略：mamba/src/play/strategy_profiles.py",
    "传射：mamba/src/tactics/targeting/attack.py",
    "运动：mamba/src/tactics/motion.py",
    "定位球：mamba/src/tactics/targeting/restart.py",
    "测试：mamba/tests/test_competition_rules.py",
    "最终包：mamba/build/com.example.mamba-2.2.0.agent",
]: bullet(doc, s)

heading(doc, "附录 B：最终验收清单", 1)
for s in [
    "[ ] Studio 显示 football3v3。", "[ ] 三台己方机器人在线且位姿持续更新。",
    "[ ] GameControl 状态持续更新且不过期。", "[ ] SET、Stop、FINISHED 完全停止。",
    "[ ] Kick-off 双方站位合规。", "[ ] 对方定位球保持至少 1.45m。",
    "[ ] 受罚机器人停止，其他机器人重新分配角色。", "[ ] 门将受罚时存在替补门将。",
    "[ ] 前场进入 Precision，后场进入 Defensive。", "[ ] 最后150秒根据比分正确切换。",
    "[ ] 无连续异常、撞场、频繁角色抖动或大比例出界。", "[ ] 提交文件名和内部版本均为2.2.0。",
]: bullet(doc, s)

doc.core_properties.title = "Mamba 3v3 完整对话与开发报告"
doc.core_properties.subject = "Booster 3v3 足球 Agent 赛事规则、开发、调参与提交归档"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
